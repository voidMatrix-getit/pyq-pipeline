"""
step3_docx.py
Converts final pipe-separated lines into a Word document table.
Uses python-docx (pip install python-docx).
"""

import logging
from pathlib import Path
from scripts.utils import parse_pipe_line, COLUMNS

log = logging.getLogger(__name__)

# Column widths in inches (total ≈ 10.5 for A4 landscape-ish)
COL_WIDTHS_CM = {
    "No":             0.6,
    "Section":        1.8,
    "Sub-Section":    2.0,
    "Question":       3.8,
    "Answer 1":       2.0,
    "Answer 2":       2.0,
    "Answer 3":       2.0,
    "Answer 4":       2.0,
    "Correct Answer": 1.2,
    "Explanation":    3.8,
    "Difficulty":     1.0,
}

HEADER_COLOR = "1F4E79"   # dark blue
ALT_ROW_COLOR = "EBF3FB"  # light blue


def cm_to_emu(cm: float) -> int:
    return int(cm * 360000)


def generate_docx(lines: list[str], output_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx not installed. Run: pip install python-docx")
        return

    doc = Document()

    # Page setup — A4 landscape
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RRB NTPC — Question Bank")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()  # spacer

    rows = []
    for line in lines:
        row = parse_pipe_line(line)
        if row:
            rows.append(row)
        else:
            log.warning(f"Skipped malformed line: {line[:60]}")

    if not rows:
        log.error("No valid rows to write to DOCX.")
        return

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(COLUMNS))
    table.style = "Table Grid"

    # Set column widths
    for j, col in enumerate(COLUMNS):
        for cell in table.columns[j].cells:
            cell.width = Cm(COL_WIDTHS_CM.get(col, 1.5))

    # Header row
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(COLUMNS):
        cell = hdr_cells[j]
        cell.text = col
        para = cell.paragraphs[0]
        run  = para.runs[0] if para.runs else para.add_run(col)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # background color
        _set_cell_bg(cell, HEADER_COLOR)

    # Data rows
    for i, row in enumerate(rows):
        tr_cells = table.rows[i + 1].cells
        for j, col in enumerate(COLUMNS):
            cell = tr_cells[j]
            cell.text = row.get(col, "")
            para = cell.paragraphs[0]
            run  = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.size = Pt(8.5)
            if col == "Correct Answer" and "UNCERTAIN" in row.get(col, ""):
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.bold = True
            # Alternate row shading
            if i % 2 == 1:
                _set_cell_bg(cell, ALT_ROW_COLOR)

    doc.save(output_path)
    log.info(f"DOCX written: {output_path} ({len(rows)} rows)")


def _set_cell_bg(cell, hex_color: str):
    """Set cell background color via XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)